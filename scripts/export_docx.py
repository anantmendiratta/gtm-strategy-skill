#!/usr/bin/env python3
"""
Export a GTM strategy markdown file to a formatted DOCX document.

Usage:
    python export_docx.py <input.md> [output.docx]

If output path is omitted, saves alongside the input file with .docx extension.

Requires: python-docx
Install:  pip install python-docx
"""

import sys
import re
from pathlib import Path


def check_dependencies():
    try:
        from docx import Document
        return True
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        return True


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx, return (rows, next_idx)."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[-| :]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows, style_header=True):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(cell_text)
            run.font.size = Pt(10)
            if r_idx == 0 and style_header:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "1B5E4B")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)

    doc.add_paragraph()


def build_docx(md_path: Path, out_path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Style helpers
    def set_heading(para, level):
        colors = {1: "1B5E4B", 2: "1B5E4B", 3: "2E7D5E"}
        sizes  = {1: 22, 2: 16, 3: 13}
        run = para.runs[0] if para.runs else para.add_run(para.text)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 12))
        run.font.color.rgb = RGBColor(
            int(colors.get(level, "000000")[:2], 16),
            int(colors.get(level, "000000")[2:4], 16),
            int(colors.get(level, "000000")[4:], 16),
        )

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            para = doc.add_heading(line[2:].strip(), level=1)
            set_heading(para, 1)

        # H2
        elif line.startswith("## ") and not line.startswith("### "):
            para = doc.add_heading(line[3:].strip(), level=2)
            set_heading(para, 2)

        # H3
        elif line.startswith("### "):
            para = doc.add_heading(line[4:].strip(), level=3)
            set_heading(para, 3)

        # Table
        elif line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        # Blockquote (investor update)
        elif line.strip().startswith(">"):
            content = line.strip().lstrip("> ").strip("*").strip()
            para = doc.add_paragraph(content)
            para.paragraph_format.left_indent = Inches(0.4)
            para.paragraph_format.right_indent = Inches(0.4)
            for run in para.runs:
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x5E)

        # Bullet
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            content = line.strip()[2:]
            # Handle bold in bullets
            para = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"\*\*(.+?)\*\*", content)
            for idx, part in enumerate(parts):
                run = para.add_run(part)
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.bold = True

        # Numbered list
        elif re.match(r"^\d+\. ", line.strip()):
            content = re.sub(r"^\d+\. ", "", line.strip())
            para = doc.add_paragraph(style="List Number")
            run = para.add_run(content)
            run.font.size = Pt(10.5)

        # Horizontal rule — skip
        elif line.strip() in ("---", "***", "___"):
            pass

        # Code block — skip fence markers, render content as monospace
        elif line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                if lines[i].strip():
                    para = doc.add_paragraph(lines[i])
                    for run in para.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                i += 1

        # Bold line (standalone)
        elif line.strip().startswith("**") and line.strip().endswith("**"):
            content = line.strip().strip("*")
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.bold = True
            run.font.size = Pt(11)

        # Regular paragraph
        elif line.strip():
            parts = re.split(r"\*\*(.+?)\*\*", line.strip())
            para = doc.add_paragraph()
            for idx, part in enumerate(parts):
                run = para.add_run(part)
                run.font.size = Pt(10.5)
                if idx % 2 == 1:
                    run.bold = True

        i += 1

    doc.save(str(out_path))
    print(f"Saved: {out_path}")


def main():
    check_dependencies()

    if len(sys.argv) < 2:
        print("Usage: python export_docx.py <input.md> [output.docx]")
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"File not found: {md_path}")
        sys.exit(1)

    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")

    build_docx(md_path, out_path)


if __name__ == "__main__":
    main()
